from io import BytesIO
from typing import Dict, List

import docxtpl
import jinja2
import pymorphy2
from docx import Document
from docx.enum.text import WD_COLOR_INDEX
from docxtpl import DocxTemplate

morph = pymorphy2.MorphAnalyzer()


class CustomFilters:
    # Вспомогательные фильтры шаблонов
    def __init__(self):
        self._enabled = True

    def enable(self, enable_filters: bool):
        """Активирует/деактивирует все пользовательские фильтры класса"""
        self._enabled = enable_filters

    def fio_short(self, fio: str) -> str:
        """Преобразует 'Фамилия Имя Отчество' к виду 'Фамилия И.О."""
        if not self._enabled:
            return fio
        fields = fio.split(" ")
        initials = "".join(field[0].capitalize() + "." for field in fields[1:])
        return " ".join([fields[0].capitalize(), initials])

    def fio_title(self, fio: str) -> str:
        """Преобразует 'фамилия имя отчество' к виду 'Фамилия Имя Отчество"""
        if not self._enabled:
            return fio
        return fio.title()

    def inflect_word(self, word: str, case: str) -> str:
        """Преобразование слова в заданный падеж

        :param case:
        'nomn' именительный, 'gent' родительный, 'datv' дательный,
        'accs' винительный, 'ablt' творительный, 'loct' предложный,
        'voct' звательный
        """
        try:
            p = next(filter(lambda x: {"nomn"} in x.tag, morph.parse(word)))
        except StopIteration:
            print("Not found nomn form for ", word)
            return word
        return p.inflect({case}).word

    def inflect_words(self, words: str, case: str) -> str:
        """Преобразование каждого из слов в строке в заданный падеж

        :param case:
        'nomn' именительный, 'gent' родительный, 'datv' дательный,
        'accs' винительный, 'ablt' творительный, 'loct' предложный,
        'voct' звательный
        """
        return " ".join(self.inflect_word(w, case) for w in words.split(" "))

    def genitive(self, words: str) -> str:
        """Преобразует слова родительный падеж"""
        if not self._enabled:
            return words
        return self.inflect_words(words, "gent")

    def dative(self, words: str) -> str:
        """Преобразует слова дательный падеж"""
        if not self._enabled:
            return words
        return self.inflect_words(words, "datv")

    def get_filters(self):
        """Возращает словарь вида {тег:функция} для всех фильтров"""
        filters = {
            "fio_short": self.fio_short,
            "fio_title": self.fio_title,
            "genitive": self.genitive,
            "dative": self.dative,
        }
        return filters


class DocumentTemplate:
    def __init__(self, template_file_name: str):
        self._template: DocxTemplate = docxtpl.DocxTemplate(template_file_name)
        self._jinja_env = jinja2.Environment()
        self._customfilters = CustomFilters()
        self._jinja_env.filters.update(self._customfilters.get_filters())

    def get_document(self, context: Dict[str, str]) -> BytesIO:
        """Генерирует и возвращает документ согласно заданному контексту"""
        self._template.render(context, jinja_env=self._jinja_env)
        file_stream = BytesIO()
        self._template.save(file_stream)
        file_stream.seek(0)
        return file_stream

    def get_draft(self, context: Dict[str, str]) -> BytesIO:
        """Генерирует и возвращает эскиз документа согласно контексту"""
        self._customfilters.enable(False)  # switch custom filters off
        self.markdown_tags()
        self._template.render(context, jinja_env=self._jinja_env)
        self._customfilters.enable(True)  # switch filters on
        file_stream = BytesIO()
        self._template.save(file_stream)
        file_stream.seek(0)
        return file_stream

    # def render(self, context: Dict[str, str]):
    #     self._template.render(context, jinja_env=self._jinja_env)

    # def save(self, result_file_name):
    #     self._template.save(result_file_name)

    def get_tags(self) -> List[str]:
        """Ищет и возвращает список тэгов подставновок переменных из шаблона"""
        return self._template.get_undeclared_template_variables(
            jinja_env=self._jinja_env
        )

    def markdown_tags(self, color=WD_COLOR_INDEX.YELLOW):
        """Размечает места тэгов заданным цветом."""
        self._template.init_docx()
        self._markdown_tag(self._template.docx, color, "{{")
        self._markdown_tag(self._template.docx, color, "}}")
        # self._template.is_rendered = True

    def _markdown_tag(self, docx: Document, color, tag: str = "{{"):
        """
        Подсветка заданного тега в документе при помощи заданного цвета.
        """

        def docx_paragraphs(docx: Document):
            """Генератор по всем параграфам документа"""
            for p in docx.paragraphs:
                yield p
            for table in docx.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for p in cell.paragraphs:
                            yield p

        for p in docx_paragraphs(docx):
            for r in p.runs:
                if tag in r.text:
                    r.font.highlight_color = color


if __name__ == "__main__":
    pass
    # Примеры использования
    # # словарь соответствия тэга из шаблона наименованию тэга (Для эсикиза)
    # tags_names = {
    #     "Наименование_организации": "Наименование организации",
    #     "дата": "дата",
    #     "Должность_руководителя": "Должность руководителя",
    #     "ФИО_руководителя": "ФИО руководителя",
    #     "ФИО": "ФИО",
    #     "продолжительность": "продолжительность",
    #     "дата_начала": "дата начала",
    #     "должность_заявителя": "должность заявителя",
    # }

    # path = "D:\\Dev\\document-template\\backend\\backend\\core\\"
    # # получение всех тэгов из шаблона
    # tpl = DocumentTemplate(path + "Заявление_на_отпуск_tpl.docx")
    # print(tpl.get_tags())

    # # подсветка всех тэгов переменных и замена их на наименования для эскиза
    # tpl = DocumentTemplate(path + "Заявление_на_отпуск_tpl.docx")
    # fs = tpl.get_draft(tags_names)
    # with open(path + "Заявление_на_отпуск_draft.docx", "wb") as outfile:
    #     # Copy the BytesIO stream to the output file
    #     outfile.write(fs.getbuffer())

    # # словарь соответствия тэга из шаблона значениям для документа
    # tags_values = {
    #     "Должность_руководителя": "генеральный директор",
    #     "Наименование_организации": '"ООО "Рога и копыта"',
    #     "ФИО_руководителя": "рогов федор федорович",
    #     "ФИО": "иванов иван иванович",
    #     "продолжительность": "14",
    #     "дата_начала": "23.10.2023",
    #     "должность_заявителя": "старший рогополировальщик",
    #     "дата": "16.10.2023",
    # }

    # # Генерация документа по заданному контексту значений полей
    # tpl = DocumentTemplate(path + "Заявление_на_отпуск_tpl.docx")
    # fs = tpl.get_document(tags_values)
    # with open(path + "Заявление_на_отпуск_doc.docx", "wb") as outfile:
    #     # Copy the BytesIO stream to the output file
    #     outfile.write(fs.getbuffer())
